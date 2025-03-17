from io import BytesIO
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from docx import Document
import base64
from aiogram import Bot

SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
]

class GoogleDocsProcessor:
    def __init__(self, credentials_path: str, template_id: str, destination_folder_id: str):
        self.creds = service_account.Credentials.from_service_account_file(
            credentials_path, scopes=SCOPES
        )
        self.drive_service = build("drive", "v3", credentials=self.creds)
        self.docs_service = build("docs", "v1", credentials=self.creds)
        self.template_id = template_id
        self.destination_folder_id = destination_folder_id

    def download_template_as_docx(self) -> BytesIO:
        request = self.drive_service.files().export_media(
            fileId=self.template_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = request.execute()
        return BytesIO(response)

    def fill_docx_with_data(self, docx_content: BytesIO, replacements: dict, table_data: list) -> BytesIO:
        document = Document(docx_content)
        keyword = "{{table}}"

        # Replace placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        # Replace placeholders in normal paragraphs
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Insert table if "{{table}}" placeholder exists
        for paragraph in document.paragraphs:
            if keyword in paragraph.text:
                paragraph.text = paragraph.text.replace(keyword, "")
                if table_data:
                    table = document.add_table(rows=1, cols=len(table_data[0]))
                    table.style = "Table Grid"  # Optional: Set table style

                    # Insert header row
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(table_data[0]):
                        hdr_cells[i].text = header

                    # Insert remaining rows
                    for row_data in table_data[1:]:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row_data):
                            row_cells[i].text = cell
                break

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

    def upload_docx_to_drive(self, docx_content: BytesIO, file_name: str) -> str:
        file_metadata = {
            "name": file_name,
            "mimeType": "application/vnd.google-apps.document",
            "parents": [self.destination_folder_id],  # Ensures file is saved in the correct folder
        }
        media = MediaIoBaseUpload(
            docx_content,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        file = (
            self.drive_service.files()
            .create(body=file_metadata, media_body=media, fields="id")
            .execute()
        )
        print("Uploading to Google Drive...")
        return file["id"]

    def download_doc_as_pdf(self, doc_id: str) -> bytes:
        request = self.drive_service.files().export_media(
            fileId=doc_id, mimeType="application/pdf"
        )
        return request.execute()

    async def process_document(self, bot: Bot, chat_id: str, replacements: dict, file_name: str, table_data: list):
        await bot.send_message(chat_id, "Template yuklab olinmoqda...")
        docx_content = self.download_template_as_docx()

        await bot.send_message(chat_id, "Ma'lumotlar shablon bilan to'ldirilmoqda...")
        filled_docx = self.fill_docx_with_data(docx_content, replacements, table_data)

        await bot.send_message(chat_id, "Hujjat Google Drive'ga yuklanmoqda...")
        new_doc_id = self.upload_docx_to_drive(filled_docx, file_name)

        await bot.send_message(chat_id, "PDF formatga o'tkazilib yuklab olinmoqda...")
        pdf_content = self.download_doc_as_pdf(new_doc_id)

        # await bot.send_message(chat_id, "PDF fayl base64 formatiga o'tkazilmoqda...")
        pdf_base64 = base64.b64encode(pdf_content).decode("utf-8")

        await bot.send_message(chat_id, "Hujjat tayyor! Endi didoxda shartnoma yaratilmoqda ...")
        return pdf_base64, pdf_content


# if __name__ == "__main__":
#     credentials_path = "service-account.json"
#     template_id = "1pRNYL7ITxcdYWfr8nO2JWSqGCYjr25fhu1lA9QzG8A8"
#     destination_folder_id = "1aUC2DkkYgmDmX3uXNAIj6KGKej5izOKn"

#     processor = GoogleDocsProcessor(credentials_path, template_id, destination_folder_id)

#     replacements = {
#         "company_name": "ZM-1",
#         "company_owner": "Komiljonov Shukurulloh Ravshan o'g'li",
#     }
#     file_name = "Filled_Template"

#     table_data = [
#         ["Column1", "Column2", "Column3"],
#         ["Column1 Row1", "Column2 Row1", "Column3 Row1"],
#         ["Column1 Row2", "Column2 Row2", "Column3 Row2"],
#     ]

#     pdf_content = processor.process_document(replacements, file_name, table_data)
    
